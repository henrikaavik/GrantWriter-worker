"""Generate final application documents."""

from typing import Dict, Any, Optional
from .gemini_client import GeminiService
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
import io


class DocumentGenerator:
    """Generate final application documents."""

    def __init__(self, language: str = "et"):
        """
        Initialize document generator.

        Args:
            language: Language for generated content ('et' or 'en')
        """
        self.gemini = GeminiService(language)
        self.language = language

    def generate_application_docx(
        self,
        project: dict,
        requirements_text: str
    ) -> Optional[bytes]:
        """
        Generate application DOCX file.

        Args:
            project: Project data with documents
            requirements_text: Grant requirements text

        Returns:
            DOCX file as bytes or None on error
        """
        # Compile document texts
        doc_texts = {}
        for doc in project.get("project_documents", []):
            if doc.get("extracted_text"):
                doc_texts[doc["name"]] = doc["extracted_text"]

        project_info = {
            "name": project.get("name", ""),
            "description": project.get("description", ""),
            "grant_name": project.get("grants", {}).get("name", "")
        }

        # Generate content sections using AI
        summary = self.gemini.generate_content(
            project_info=project_info,
            documents_text=doc_texts,
            requirements_text=requirements_text,
            content_type="summary"
        )

        narrative = self.gemini.generate_content(
            project_info=project_info,
            documents_text=doc_texts,
            requirements_text=requirements_text,
            content_type="narrative"
        )

        if not summary and not narrative:
            return None

        # Create DOCX document
        doc = Document()

        # Title
        title = doc.add_heading(project_info["name"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle with grant name
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(
            f"Application for: {project_info['grant_name']}"
        )
        subtitle_run.italic = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # Executive Summary
        if summary:
            doc.add_heading(
                "Executive Summary" if self.language == "en" else "Kokkuvõte",
                level=1
            )
            doc.add_paragraph(summary)

        # Project Narrative
        if narrative:
            doc.add_heading(
                "Project Narrative" if self.language == "en" else "Projekti kirjeldus",
                level=1
            )

            # Split narrative into paragraphs
            for para in narrative.split("\n\n"):
                if para.strip():
                    # Check if it's a heading (starts with # or **)
                    if para.strip().startswith("#"):
                        heading_text = para.strip().lstrip("#").strip()
                        doc.add_heading(heading_text, level=2)
                    elif para.strip().startswith("**") and para.strip().endswith("**"):
                        heading_text = para.strip().strip("*").strip()
                        doc.add_heading(heading_text, level=2)
                    else:
                        doc.add_paragraph(para.strip())

        # Add section for uploaded documents reference
        doc.add_heading(
            "Supporting Documents" if self.language == "en" else "Lisadokumendid",
            level=1
        )

        doc.add_paragraph(
            "The following documents are included with this application:"
            if self.language == "en"
            else "Taotlusega on kaasatud järgmised dokumendid:"
        )

        for doc_name in doc_texts.keys():
            doc.add_paragraph(f"• {doc_name}", style="List Bullet")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def generate_budget_xlsx(
        self,
        project: dict,
        requirements_text: str
    ) -> Optional[bytes]:
        """
        Generate budget XLSX file.

        Args:
            project: Project data with documents
            requirements_text: Grant requirements text

        Returns:
            XLSX file as bytes or None on error
        """
        # Compile document texts
        doc_texts = {}
        for doc in project.get("project_documents", []):
            if doc.get("extracted_text"):
                doc_texts[doc["name"]] = doc["extracted_text"]

        project_info = {
            "name": project.get("name", ""),
            "description": project.get("description", "")
        }

        # Generate budget content using AI
        budget_content = self.gemini.generate_content(
            project_info=project_info,
            documents_text=doc_texts,
            requirements_text=requirements_text,
            content_type="budget"
        )

        # Create XLSX workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Budget" if self.language == "en" else "Eelarve"

        # Styles
        header_font = Font(bold=True, size=12)
        header_alignment = Alignment(horizontal="center", vertical="center")
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin")
        )

        # Title
        ws.merge_cells("A1:D1")
        ws["A1"] = f"Budget - {project_info['name']}"
        ws["A1"].font = Font(bold=True, size=14)
        ws["A1"].alignment = Alignment(horizontal="center")

        # Headers
        headers = ["Category", "Description", "Amount (EUR)", "Justification"]
        if self.language == "et":
            headers = ["Kategooria", "Kirjeldus", "Summa (EUR)", "Põhjendus"]

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=header)
            cell.font = header_font
            cell.alignment = header_alignment
            cell.border = thin_border

        # Parse budget content and add rows
        row = 4
        total = 0

        if budget_content:
            # Try to parse structured content
            lines = budget_content.split("\n")
            for line in lines:
                line = line.strip()
                if not line or line.startswith("#") or line.startswith("---"):
                    continue

                # Try to parse table row (pipe-separated)
                if "|" in line:
                    parts = [p.strip() for p in line.split("|") if p.strip()]
                    if len(parts) >= 3 and not all(c in "-|" for c in line):
                        category = parts[0] if len(parts) > 0 else ""
                        description = parts[1] if len(parts) > 1 else ""
                        amount_str = parts[2] if len(parts) > 2 else "0"
                        justification = parts[3] if len(parts) > 3 else ""

                        # Try to parse amount
                        try:
                            amount = float(
                                amount_str.replace(",", "")
                                .replace("€", "")
                                .replace("EUR", "")
                                .strip()
                            )
                        except ValueError:
                            amount = 0

                        if category and category.lower() not in ["category", "kategooria"]:
                            ws.cell(row=row, column=1, value=category).border = thin_border
                            ws.cell(row=row, column=2, value=description).border = thin_border
                            ws.cell(row=row, column=3, value=amount).border = thin_border
                            ws.cell(row=row, column=4, value=justification).border = thin_border
                            total += amount
                            row += 1

        # If no structured content, add placeholder rows
        if row == 4:
            default_categories = [
                ("Personnel", "Staff costs", 0, ""),
                ("Equipment", "Equipment and materials", 0, ""),
                ("Travel", "Travel and meetings", 0, ""),
                ("Other", "Other direct costs", 0, ""),
                ("Overhead", "Indirect costs", 0, "")
            ]
            if self.language == "et":
                default_categories = [
                    ("Personal", "Tööjõukulud", 0, ""),
                    ("Seadmed", "Seadmed ja materjalid", 0, ""),
                    ("Reisid", "Reisi- ja koosolekukulud", 0, ""),
                    ("Muud", "Muud otsesed kulud", 0, ""),
                    ("Üldkulud", "Kaudsed kulud", 0, "")
                ]

            for cat, desc, amt, just in default_categories:
                ws.cell(row=row, column=1, value=cat).border = thin_border
                ws.cell(row=row, column=2, value=desc).border = thin_border
                ws.cell(row=row, column=3, value=amt).border = thin_border
                ws.cell(row=row, column=4, value=just).border = thin_border
                row += 1

        # Total row
        row += 1
        ws.cell(row=row, column=1, value="TOTAL" if self.language == "en" else "KOKKU")
        ws.cell(row=row, column=1).font = header_font
        ws.cell(row=row, column=3, value=total)
        ws.cell(row=row, column=3).font = header_font

        # Adjust column widths
        ws.column_dimensions["A"].width = 20
        ws.column_dimensions["B"].width = 40
        ws.column_dimensions["C"].width = 15
        ws.column_dimensions["D"].width = 40

        # Save to bytes
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def generate_docx_from_sections(
        self,
        project_name: str,
        grant_name: str,
        sections: list
    ) -> Optional[bytes]:
        """
        Generate final DOCX from pre-edited sections.

        Args:
            project_name: Name of the project
            grant_name: Name of the grant
            sections: List of section dicts with 'section_name' and 'content'

        Returns:
            DOCX file as bytes or None on error
        """
        try:
            doc = Document()

            # Title
            title = doc.add_heading(project_name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Subtitle with grant name
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(
                f"Taotlus: {grant_name}" if self.language == "et"
                else f"Application for: {grant_name}"
            )
            subtitle_run.italic = True
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacing

            # Add each section
            for section in sections:
                section_name = section.get("section_name", "")
                content = section.get("content", "")

                if not section_name:
                    continue

                # Add section heading
                doc.add_heading(section_name, level=1)

                if content:
                    # Split content into paragraphs and process
                    for para in content.split("\n\n"):
                        para = para.strip()
                        if not para:
                            continue

                        # Check if it's a subheading (starts with ## or **)
                        if para.startswith("##"):
                            heading_text = para.lstrip("#").strip()
                            doc.add_heading(heading_text, level=2)
                        elif para.startswith("**") and para.endswith("**"):
                            heading_text = para.strip("*").strip()
                            doc.add_heading(heading_text, level=2)
                        elif para.startswith("- ") or para.startswith("* "):
                            # Handle bullet points
                            lines = para.split("\n")
                            for line in lines:
                                line = line.strip()
                                if line.startswith("- ") or line.startswith("* "):
                                    bullet_text = line[2:].strip()
                                    doc.add_paragraph(bullet_text, style="List Bullet")
                                elif line:
                                    doc.add_paragraph(line)
                        else:
                            # Regular paragraph - handle single newlines as line breaks
                            paragraph = doc.add_paragraph()
                            lines = para.split("\n")
                            for i, line in enumerate(lines):
                                if i > 0:
                                    paragraph.add_run("\n")
                                paragraph.add_run(line.strip())
                else:
                    # Empty section placeholder
                    empty_para = doc.add_paragraph()
                    empty_run = empty_para.add_run(
                        "[Sisu puudub]" if self.language == "et"
                        else "[No content]"
                    )
                    empty_run.italic = True

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return buffer.getvalue()
        except Exception as e:
            print(f"Error generating DOCX from sections: {e}")
            return None

    def generate_cover_letter_docx(
        self,
        project: dict,
        requirements_text: str
    ) -> Optional[bytes]:
        """
        Generate a formal cover letter DOCX.

        Args:
            project: Project data with documents
            requirements_text: Grant requirements text

        Returns:
            DOCX file as bytes or None on error
        """
        try:
            # Compile document texts
            doc_texts = {}
            for doc in project.get("project_documents", []):
                if doc.get("extracted_text"):
                    doc_texts[doc["name"]] = doc["extracted_text"]

            project_info = {
                "name": project.get("name", ""),
                "description": project.get("description", ""),
                "grant_name": project.get("grants", {}).get("name", "")
            }

            # Generate cover letter content using AI
            content = self.gemini.generate_content(
                project_info=project_info,
                documents_text=doc_texts,
                requirements_text=requirements_text,
                content_type="cover_letter"
            )

            if not content:
                return None

            # Create DOCX document
            doc = Document()

            # Title
            title = doc.add_heading(
                "Kaaskiri" if self.language == "et" else "Cover Letter",
                0
            )
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacing

            # Add the cover letter content
            for para in content.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return buffer.getvalue()
        except Exception as e:
            print(f"Error generating cover letter: {e}")
            return None

    def generate_executive_summary_docx(
        self,
        project: dict,
        requirements_text: str
    ) -> Optional[bytes]:
        """
        Generate an executive summary DOCX.

        Args:
            project: Project data with documents
            requirements_text: Grant requirements text

        Returns:
            DOCX file as bytes or None on error
        """
        try:
            # Compile document texts
            doc_texts = {}
            for doc in project.get("project_documents", []):
                if doc.get("extracted_text"):
                    doc_texts[doc["name"]] = doc["extracted_text"]

            project_info = {
                "name": project.get("name", ""),
                "description": project.get("description", ""),
                "grant_name": project.get("grants", {}).get("name", "")
            }

            # Generate executive summary content using AI
            content = self.gemini.generate_content(
                project_info=project_info,
                documents_text=doc_texts,
                requirements_text=requirements_text,
                content_type="executive_summary"
            )

            if not content:
                return None

            # Create DOCX document
            doc = Document()

            # Title
            title = doc.add_heading(
                "Kokkuvõte" if self.language == "et" else "Executive Summary",
                0
            )
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Subtitle with project name
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(project_info["name"])
            subtitle_run.italic = True
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacing

            # Add the summary content with heading support
            for para in content.split("\n\n"):
                para = para.strip()
                if not para:
                    continue

                # Check if it's a heading
                if para.startswith("##"):
                    heading_text = para.lstrip("#").strip()
                    doc.add_heading(heading_text, level=2)
                elif para.startswith("#"):
                    heading_text = para.lstrip("#").strip()
                    doc.add_heading(heading_text, level=1)
                elif para.startswith("**") and para.endswith("**"):
                    heading_text = para.strip("*").strip()
                    doc.add_heading(heading_text, level=2)
                else:
                    doc.add_paragraph(para)

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return buffer.getvalue()
        except Exception as e:
            print(f"Error generating executive summary: {e}")
            return None

    def generate_timeline_xlsx(
        self,
        project: dict,
        requirements_text: str
    ) -> Optional[bytes]:
        """
        Generate a project timeline XLSX.

        Args:
            project: Project data with documents
            requirements_text: Grant requirements text

        Returns:
            XLSX file as bytes or None on error
        """
        try:
            # Compile document texts
            doc_texts = {}
            for doc in project.get("project_documents", []):
                if doc.get("extracted_text"):
                    doc_texts[doc["name"]] = doc["extracted_text"]

            project_info = {
                "name": project.get("name", ""),
                "description": project.get("description", "")
            }

            # Generate timeline content using AI
            timeline_content = self.gemini.generate_content(
                project_info=project_info,
                documents_text=doc_texts,
                requirements_text=requirements_text,
                content_type="timeline"
            )

            # Create XLSX workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "Ajakava" if self.language == "et" else "Timeline"

            # Styles
            header_font = Font(bold=True, size=12)
            header_alignment = Alignment(horizontal="center", vertical="center")
            thin_border = Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin")
            )

            # Title
            ws.merge_cells("A1:E1")
            ws["A1"] = f"{'Ajakava' if self.language == 'et' else 'Timeline'} - {project_info['name']}"
            ws["A1"].font = Font(bold=True, size=14)
            ws["A1"].alignment = Alignment(horizontal="center")

            # Headers
            headers = ["Phase", "Activities", "Start", "End", "Deliverables"]
            if self.language == "et":
                headers = ["Faas", "Tegevused", "Algus", "Lõpp", "Tulemid"]

            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=3, column=col, value=header)
                cell.font = header_font
                cell.alignment = header_alignment
                cell.border = thin_border

            # Parse timeline content and add rows
            row = 4

            if timeline_content:
                lines = timeline_content.split("\n")
                for line in lines:
                    line = line.strip()
                    if not line or line.startswith("#") or line.startswith("---"):
                        continue

                    # Try to parse table row (pipe-separated)
                    if "|" in line:
                        parts = [p.strip() for p in line.split("|") if p.strip()]
                        if len(parts) >= 3 and not all(c in "-|" for c in line):
                            phase = parts[0] if len(parts) > 0 else ""
                            activities = parts[1] if len(parts) > 1 else ""
                            start = parts[2] if len(parts) > 2 else ""
                            end = parts[3] if len(parts) > 3 else ""
                            deliverables = parts[4] if len(parts) > 4 else ""

                            # Skip header row
                            if phase.lower() in ["phase", "faas"]:
                                continue

                            ws.cell(row=row, column=1, value=phase).border = thin_border
                            ws.cell(row=row, column=2, value=activities).border = thin_border
                            ws.cell(row=row, column=3, value=start).border = thin_border
                            ws.cell(row=row, column=4, value=end).border = thin_border
                            ws.cell(row=row, column=5, value=deliverables).border = thin_border
                            row += 1

            # If no structured content, add placeholder rows
            if row == 4:
                default_phases = [
                    ("Initiation", "Project setup, team formation", "Month 1", "Month 1", "Project plan"),
                    ("Development", "Core development work", "Month 2", "Month 6", "Prototype"),
                    ("Testing", "Testing and validation", "Month 6", "Month 8", "Test results"),
                    ("Implementation", "Deployment and rollout", "Month 8", "Month 10", "Deployed system"),
                    ("Closeout", "Final reporting", "Month 10", "Month 12", "Final report")
                ]
                if self.language == "et":
                    default_phases = [
                        ("Alustamine", "Projekti seadistamine, meeskonna moodustamine", "Kuu 1", "Kuu 1", "Projektiplaan"),
                        ("Arendus", "Põhiarendus", "Kuu 2", "Kuu 6", "Prototüüp"),
                        ("Testimine", "Testimine ja valideerimine", "Kuu 6", "Kuu 8", "Testitulemused"),
                        ("Rakendamine", "Juurutamine", "Kuu 8", "Kuu 10", "Juurutatud süsteem"),
                        ("Lõpetamine", "Lõpparuandlus", "Kuu 10", "Kuu 12", "Lõpparuanne")
                    ]

                for phase, activities, start, end, deliverables in default_phases:
                    ws.cell(row=row, column=1, value=phase).border = thin_border
                    ws.cell(row=row, column=2, value=activities).border = thin_border
                    ws.cell(row=row, column=3, value=start).border = thin_border
                    ws.cell(row=row, column=4, value=end).border = thin_border
                    ws.cell(row=row, column=5, value=deliverables).border = thin_border
                    row += 1

            # Adjust column widths
            ws.column_dimensions["A"].width = 20
            ws.column_dimensions["B"].width = 40
            ws.column_dimensions["C"].width = 12
            ws.column_dimensions["D"].width = 12
            ws.column_dimensions["E"].width = 30

            # Save to bytes
            buffer = io.BytesIO()
            wb.save(buffer)
            buffer.seek(0)

            return buffer.getvalue()
        except Exception as e:
            print(f"Error generating timeline: {e}")
            return None

    def generate_risk_analysis_docx(
        self,
        project: dict,
        requirements_text: str
    ) -> Optional[bytes]:
        """
        Generate a risk analysis DOCX.

        Args:
            project: Project data with documents
            requirements_text: Grant requirements text

        Returns:
            DOCX file as bytes or None on error
        """
        try:
            # Compile document texts
            doc_texts = {}
            for doc in project.get("project_documents", []):
                if doc.get("extracted_text"):
                    doc_texts[doc["name"]] = doc["extracted_text"]

            project_info = {
                "name": project.get("name", ""),
                "description": project.get("description", ""),
                "grant_name": project.get("grants", {}).get("name", "")
            }

            # Generate risk analysis content using AI
            content = self.gemini.generate_content(
                project_info=project_info,
                documents_text=doc_texts,
                requirements_text=requirements_text,
                content_type="risk_analysis"
            )

            if not content:
                return None

            # Create DOCX document
            doc = Document()

            # Title
            title = doc.add_heading(
                "Riskianalüüs" if self.language == "et" else "Risk Analysis",
                0
            )
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Subtitle with project name
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(project_info["name"])
            subtitle_run.italic = True
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacing

            # Add the risk analysis content with heading support
            for para in content.split("\n\n"):
                para = para.strip()
                if not para:
                    continue

                # Check if it's a heading
                if para.startswith("##"):
                    heading_text = para.lstrip("#").strip()
                    doc.add_heading(heading_text, level=2)
                elif para.startswith("#"):
                    heading_text = para.lstrip("#").strip()
                    doc.add_heading(heading_text, level=1)
                elif para.startswith("**") and para.endswith("**"):
                    heading_text = para.strip("*").strip()
                    doc.add_heading(heading_text, level=2)
                elif para.startswith("- ") or para.startswith("* "):
                    # Handle bullet points
                    lines = para.split("\n")
                    for line in lines:
                        line = line.strip()
                        if line.startswith("- ") or line.startswith("* "):
                            bullet_text = line[2:].strip()
                            doc.add_paragraph(bullet_text, style="List Bullet")
                        elif line:
                            doc.add_paragraph(line)
                else:
                    doc.add_paragraph(para)

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return buffer.getvalue()
        except Exception as e:
            print(f"Error generating risk analysis: {e}")
            return None
