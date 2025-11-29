"""Document parsing using Docling library."""

from typing import Dict, Any, Optional, List
import tempfile
import os


class DocumentParser:
    """Parse documents using Docling library."""

    def __init__(self):
        """Initialize the document parser."""
        self._converter = None

    def _get_converter(self):
        """Lazy load the document converter."""
        if self._converter is None:
            try:
                from docling.document_converter import DocumentConverter
                self._converter = DocumentConverter()
            except ImportError:
                raise ImportError(
                    "Docling is not installed. Install with: pip install docling"
                )
        return self._converter

    def parse_document(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse a document and extract text content.

        Args:
            file_bytes: Raw file bytes
            filename: Original filename for extension detection

        Returns:
            Dict with 'text', 'markdown', 'metadata'
        """
        # Get file extension
        suffix = os.path.splitext(filename)[1].lower()

        # For PDFs, try lightweight extraction first to save memory
        if suffix == ".pdf":
            text = self._try_pypdf2(file_bytes)
            if text and len(text.strip()) > 100:  # Got meaningful text
                print(f"Using PyPDF2 (fast) for {filename}")
                return {"text": text, "markdown": text, "metadata": {"method": "pypdf2"}}
            print(f"PyPDF2 insufficient, trying docling/OCR for {filename}")

        # Write to temp file (Docling needs file path)
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            converter = self._get_converter()
            result = converter.convert(tmp_path)

            # Extract content
            text = ""
            markdown = ""
            metadata = {}

            try:
                text = result.document.export_to_text()
            except Exception:
                pass

            try:
                markdown = result.document.export_to_markdown()
            except Exception:
                pass

            try:
                # Get document metadata
                if hasattr(result.document, 'pages'):
                    metadata['page_count'] = len(result.document.pages)

                # Count tables if available
                if hasattr(result.document, 'tables'):
                    metadata['table_count'] = len(result.document.tables)
            except Exception:
                pass

            return {
                "text": text,
                "markdown": markdown,
                "metadata": metadata
            }

        except Exception as e:
            # Fallback: try basic text extraction
            return self._fallback_extract(file_bytes, filename, str(e))

        finally:
            # Clean up temp file
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    def _try_pypdf2(self, file_bytes: bytes) -> str:
        """Try lightweight PDF extraction with PyPDF2."""
        try:
            import io
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            text = "\n".join(
                page.extract_text() or ""
                for page in reader.pages
            )
            return text
        except Exception as e:
            print(f"PyPDF2 extraction failed: {e}")
            return ""

    def _fallback_extract(
        self,
        file_bytes: bytes,
        filename: str,
        error_msg: str
    ) -> Dict[str, Any]:
        """
        Fallback text extraction when Docling fails.

        Args:
            file_bytes: Raw file bytes
            filename: Original filename
            error_msg: Original error message

        Returns:
            Dict with extracted text or error info
        """
        ext = os.path.splitext(filename)[1].lower()
        text = ""

        try:
            if ext == ".txt":
                text = file_bytes.decode("utf-8", errors="ignore")

            elif ext == ".pdf":
                # Try PyPDF2 as fallback
                try:
                    import io
                    from PyPDF2 import PdfReader
                    reader = PdfReader(io.BytesIO(file_bytes))
                    text = "\n".join(
                        page.extract_text() or ""
                        for page in reader.pages
                    )
                except ImportError:
                    pass

            elif ext in [".docx", ".doc"]:
                # Try python-docx as fallback
                try:
                    import io
                    from docx import Document
                    doc = Document(io.BytesIO(file_bytes))
                    text = "\n".join(
                        para.text
                        for para in doc.paragraphs
                    )
                except ImportError:
                    pass

        except Exception:
            pass

        return {
            "text": text,
            "markdown": text,
            "metadata": {
                "fallback": True,
                "original_error": error_msg
            }
        }

    def extract_tables(self, file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Extract tables from a document.

        Args:
            file_bytes: Raw file bytes
            filename: Original filename

        Returns:
            List of table data
        """
        tables = []

        suffix = os.path.splitext(filename)[1].lower()

        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            converter = self._get_converter()
            result = converter.convert(tmp_path)

            # Extract tables from document
            if hasattr(result.document, 'tables'):
                for table in result.document.tables:
                    try:
                        table_data = {
                            "rows": [],
                            "headers": []
                        }
                        # Extract table structure
                        if hasattr(table, 'export_to_dataframe'):
                            df = table.export_to_dataframe()
                            table_data["headers"] = list(df.columns)
                            table_data["rows"] = df.values.tolist()
                        tables.append(table_data)
                    except Exception:
                        pass

        except Exception as e:
            print(f"Error extracting tables: {e}")

        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

        return tables


# Convenience function for simple text extraction
_parser_instance = None


def parse_document(file_bytes: bytes, filename: str) -> str:
    """
    Parse a document and return extracted text.

    Args:
        file_bytes: Raw file bytes
        filename: Original filename

    Returns:
        Extracted text content
    """
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = DocumentParser()

    result = _parser_instance.parse_document(file_bytes, filename)
    return result.get("text", "") or result.get("markdown", "")
